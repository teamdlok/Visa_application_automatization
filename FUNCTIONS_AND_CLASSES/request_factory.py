from docx import Document
from datetime import datetime
from FUNCTIONS_AND_CLASSES.FUNCTIONS import *


# Параметры стилей для документа заявлений
font_params = {
    "fontname": "Arial Narrow",
    "fontsize": 11,
    "bold": True,
}

# Таблица 0 багованная, потому для неё нужен отдельный функционал
def change_row_durdom(tables, table_number, row_number, from_indx, to_indx, text, font_params, igor_toxic = False):
    text = list(text)
    text_is_bigger = False
    counter = 0
    
    row = tables[table_number].rows[row_number]

    for index, cell in enumerate(row.cells):
        try:
            if not index % 2 == 0:

                if to_indx - from_indx + 1 < len(text):
                    text_is_bigger = True
                if index >= from_indx and index <= to_indx:
                    if igor_toxic:
                        cell.text = text
                    else:
                        cell.text = text.pop(0)
                    counter = counter + 1
                    # print(f" символ - {cell.text} счетчик - {counter}")

                    change_cell_font(cell, **font_params)
                    paragraphs = cell.paragraphs
                    paragraphs
                    for paragraph in paragraphs:
                        if "align" in font_params:
                            if font_params['align'] == "LEFT":
                                paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
                
                elif to_indx - from_indx + 1 < len(text): # Тут и скрывается проблема. не подпадая под первое условие, оно подпадает под это, и меняет строку
                    print(f"{to_indx - from_indx} меньше чем {len(text)}")

        except IndexError:
            cell.text = ""
            # print("<<ERROR>>")
            pass
        
    if text_is_bigger:
        to_index_len = len(tables[table_number].rows[row_number].cells)
        if len(tables[table_number].rows) > 1:
            change_row(tables=tables, table_number=table_number, row_number=row_number+1,from_indx=0, to_indx=to_index_len, text=text, font_params=font_params)
        else:
            change_row(tables=tables, table_number=table_number+1, row_number=row_number,from_indx=0, to_indx=to_index_len, text=text, font_params=font_params)



class ZayavlenieConstructor():
    """"Конструктор для создания заявления, получает информацию о человеке, и вписывает
    в документ для подачи заявления, работает как для новеньких так и для старых типов заявлений"""

    def __init__(self, newby_request, company_name, rnp_input, day_rnp_start, month_rnp_start,
                year_rnp_start, day_rnp_end, month_rnp_end, year_rnp_end, lastname_input,
                name_input, namechange_input,
                birthplace_input, birth_day, birth_month, month_number, birth_year, sex,
                passport_series, passport_number, passport_date_day,
                passport_date_month, passport_date_year, passport_creator,
                address_of_living, profession_name,
                ):
        super().__init__()
        self.newby_request = newby_request
        self.company_name = company_name
        self.rnp = rnp_input
        self.day_rnp_start = day_rnp_start
        self.month_rnp_start = month_rnp_start
        self.year_rnp_start = year_rnp_start
        self.day_rnp_end = day_rnp_end
        self.month_rnp_end = month_rnp_end
        self.year_rnp_end = year_rnp_end
        self.lastname = lastname_input
        self.name = name_input
        self.namechange = namechange_input
        self.birthplace = birthplace_input
        self.birth_day = birth_day
        self.birth_month = birth_month
        self.month_number = month_number
        self.birth_year = birth_year
        self.sex = sex
        self.passport_series = passport_series
        self.passport_number = passport_number
        self.passport_date_day = passport_date_day
        self.passport_date_month = passport_date_month
        self.passport_date_year = passport_date_year
        self.passport_creator = passport_creator
        self.address_of_living = address_of_living
        self.profession_name = profession_name


    def save_document(self, request_document, company_name, name, lastname ):
        current_time = datetime.now().strftime('%d.%m.%Y')
        request_document.save(f'OUTPUT/{company_name}/ЗАЯВЛЕНИЕ/{''.join(self.lastname)} {''.join(self.name)} - {current_time} КЛАССОВАЯ ВЕРСИЯ.docx')


    def request_factory(self):
        zayavlenie_path = f"./templates/{self.company_name}/ЗАЯВЛЕНИЕ/zayavlenie_template.docx"
        request_document = Document(zayavlenie_path) 
        tables = request_document.tables
        
        # Тут обрабатываем номер рнп, вносим в документ
        if self.rnp == []:
            self.rnp = "РНП № 2400"
        elif len(self.rnp) > 6:
            self.rnp = f"РНП № {''.join(self.rnp)}"
        else:
            self.rnp = f"РНП № {''.join(self.rnp)}"
            print(f"_________Ошибка? {''.join(self.rnp)}?")
        change_row_durdom(tables, 0,0,0,49, self.rnp, dict(fontname = "Times New Roman" , fontsize = 12, bold = True, align = "LEFT"), igor_toxic=True)
        
        if self.newby_request == False:
        #Тут вносится в документ заявление день начала и сверху документа и снизу
            change_row_durdom(tables, 0, 2, 4, 7, self.day_rnp_start, font_params)
            change_row(tables, 38,0,1,2, self.day_rnp_start, font_params)

            #Тут вносится в документ заявление месяц начала и сверху документа и снизу
            change_row_durdom(tables, 0, 2, 10, 13, self.month_rnp_start, font_params)
            change_row(tables, 38,0,4,5, self.month_rnp_start, font_params)
        
        change_row_durdom(tables, 0, 2, 16, 23, self.year_rnp_start, font_params)
        change_row(tables, 38,0,7,10, self.year_rnp_start, font_params)
        
        change_row_durdom(tables, 0, 2, 27, 28, self.day_rnp_end[1], font_params)
        change_row(tables, 0,2,26,26, self.day_rnp_end[0], font_params)
        change_row(tables, 38,0,12,13, self.day_rnp_end, font_params)

        change_row_durdom(tables, 0, 2, 31, 33, self.month_rnp_end, font_params)
        change_row(tables, 38,0,15,16, self.month_rnp_end, font_params)

        change_row_durdom(tables, 0, 2, 37, 43, self.year_rnp_end, font_params)
        change_row(tables, 38,0,18,21, self.year_rnp_end, font_params)

        change_row_durdom(tables, 0, 4, 6, 55, self.lastname, font_params)

        change_row(tables, 1,0,5,20, self.name, font_params)

        # Тут изменяется поле изменения имени, если оно не менялось, то если это женщина
        # То на конце добавляется не менял(а)
        if str(''.join(self.namechange)) == "НЕ МЕНЯЛ":
            namechange_check = False
        else:
            namechange_check = True
        change_row(tables, 3,0,1,24, self.namechange, font_params)

        change_row(tables, 8,0,15,28, self.birthplace, font_params)

        change_row(tables, 11,0,1,2, self.birth_day, font_params)

        change_row(tables, 11,0,4,5, self.month_number, font_params)

        change_row(tables, 11,0,7,10, self.birth_year, font_params)

        if str(''.join(self.sex)) == "М":
            change_row(tables, 11,0, 12,12, self.sex, font_params)
            change_row(tables, 11,0, 14,14, "", font_params)
        elif str(''.join(self.sex)) == "Ж":
            change_row(tables, 11,0, 12,12, "", font_params)
            change_row(tables, 11,0, 14,14, self.sex, font_params)
            if not namechange_check:
                print("В конце не менял а должно стоять")
                change_row(tables, 3,0,9,9, "А", font_params)
            else:
                print("Если указана Ж, то что-то не так")
                print(namechange_check)
                print(self.namechange)
                
            
        if len(self.passport_series) == 1:
            change_row(tables, 13, 0, 6, 6, " ", font_params)
            change_row(tables, 13, 0, 7, 7, self.passport_series, font_params)
        elif len(self.passport_series) == 2:
            change_row(tables, 13, 0, 6, 7, self.passport_series, font_params)

        if len(self.passport_number) < 6:
            print(f"_____Ошибка? В номере паспорта всего {len(self.passport_number)} цифры")
        change_row(tables, 13,0,9,16, self.passport_number, font_params)

        change_row(tables, 13,0,18,19, self.passport_date_day, font_params)

        change_row(tables, 13,0,21,22, self.passport_date_month, font_params)

        change_row(tables, 13,0,24,27, self.passport_date_year, font_params)

        change_row(tables, 14,0,1,31, self.passport_creator, font_params)

        change_row(tables, 17,0,1,21, self.address_of_living.replace(",", ""), font_params)

        change_row(tables, 20,0,1,22, self.profession_name, font_params)

        self.save_document(request_document, self.company_name, self.name, self.lastname)
