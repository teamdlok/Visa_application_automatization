import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import calendar
from docx import Document

# wordDoc = Document('C:/Users/ekole/Desktop/test/organizations/AVANTA/REQUEST_template.docx')

# tables = wordDoc.tables

# list_for_tables = list()

# for indx_of_table, table in enumerate(tables):
#     list_for_tables.append([indx_of_table])
#     for row in table.rows:
#         for index_of_cell, cell in enumerate(row.cells):
#             list_for_tables[indx_of_table].append(cell.text)
# np_list = np.array(list_for_tables, dtype="object")
# print(np_list)
# np.savetxt("tables_list_avanta_request", np_list, newline=" \n", fmt='%s')


month_dict = {}

def change_cell_font(cell, fontname, fontsize, bold = False, align="CENTER",):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        paragraph_format = paragraph
        if align == "CENTER":
            paragraph.aligment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "LEFT":
            paragraph.aligment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            raise Exception("ОШИБКА ВЫРАВНИВАНИЯ")
        for run in paragraph.runs:
            font = run.font
            font.size= Pt(fontsize)
            font.name = fontname
            
            if bold:
                font.bold = True


def change_row(tables, table_number, row_number, from_indx, to_indx, text, font_params):
    text = list(text)
    text_is_bigger = False
    counter = 0
    
    row = tables[table_number].rows[row_number]
    for index, cell in enumerate(row.cells):
        try:
            
            if to_indx - from_indx + 1 < len(text):
                text_is_bigger = True
            
            if index >= from_indx and index <= to_indx:
                cell.text = text.pop(0)
                counter = counter + 1
                # self.change_cell_font(cell, **font_params)
                change_cell_font(cell, **font_params)
                paragraphs = cell.paragraphs
                paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
            
            elif to_indx - from_indx + 1 < len(text): # Тут и скрывается проблема. не подпадая под первое условие, оно подпадает под это, и меняет строку
                print(f"{to_indx - from_indx} меньше чем {len(text)}")
        except IndexError:
            cell.text = ""
            pass
        
    if text_is_bigger:
        if len(tables[table_number].rows) > 1:
            to_index_len = len(tables[table_number].rows[row_number+1].cells)
            change_row(tables=tables, table_number=table_number, row_number=row_number+1,from_indx=0, to_indx=to_index_len, text=text, font_params=font_params)
        else:
            to_index_len = len(tables[table_number+1].rows[row_number].cells)
            change_row(tables=tables, table_number=table_number+1, row_number=row_number,from_indx=0, to_indx=to_index_len, text=text, font_params=font_params)
                

def input_prettier(field_name, normalize_digit = False, normalize_year = False, input_as_list = True):
    field_text = f"ВВЕДИТЕ {field_name} \t"
    input_field = input(field_text)
    if normalize_digit == True and not input_field == "":
        try:
            digit_prettier = int(input_field)
            if digit_prettier < 10:
                input_field = f"0{digit_prettier}"
            else:
                input_field = str(digit_prettier)
        except ValueError:
            print(f"Значение {input_field} не может быть числом")
    if normalize_year == True and not input_field == "":
        try:
            year = int(input_field)
            if year < 10:
                input_field = f"200{year}"
            elif year >= 10 and year < 40:
                input_field = f"20{year}"
            elif year > 41 and year < 100:
                input_field = f"19{year}"
            else: 
                pass
        except ValueError:
            print(f"Значение {input_field} не может быть числом")

    if input_as_list == True:   
        input_field = list(input_field.upper())
    return input_field


def prettier_check_iter(mylist):
    for i in mylist:
        yield i
        yield "_"


def prettier_check_tool(input_text):
    return ''.join(list(prettier_check_iter(input_text))[:-1])


def numbered_month(month_text):
    print(f"месяц текстом - {month_text}")
    abbr_to_num = {name: num for num, name in enumerate(calendar.month_abbr) if num}
    print(f"месяц номером - {abbr_to_num[month_text.lower().capitalize()]}")
    return abbr_to_num[month_text.lower().capitalize()]


class FiringPeople():
    """"Конструктор для удаления людей, получает информацию о человеке, и вписывает
    в документ для удаления"""

    def __init__(self, document, font_settings, lastname, name, birth_day, birth_month, birth_year, passport_series,
                 passport_number, passport_date_day, passport_date_month, passport_date_year, profession_name
                ):
        super().__init__()
        self.document = document
        self.lastname = lastname
        self.name = name
        self.birth_day = birth_day
        self.birth_month = birth_month
        self.birth_year = birth_year
        self.passport_series = passport_series
        self.passport_number = passport_number
        self.passport_date_day = passport_date_day
        self.passport_date_month = passport_date_month
        self.passport_date_year = passport_date_year
        self.profession_name = profession_name
        self.font_settings = font_settings

    def save_document(self, document, name, lastname ):
        current_time = datetime.now().strftime('%d.%m.%Y')
        document.save(f'C:/Users/ekole/Desktop/test/fired_people/{''.join(name)} {''.join(lastname)} - {current_time}.docx')


    def firing_factory(self):
        
        tables = self.document.tables
        
        change_row(tables, 22, 0, 1, 10, self.lastname, self.font_settings)
        print("ФАМИЛИЯ - ", prettier_check_tool(self.lastname))

        change_row(tables, 23,0,1,15, self.name, self.font_settings)
        print("ИМЯ - ",prettier_check_tool(self.name))

        change_row(tables, 28,0,1,2, self.birth_day, self.font_settings)
        print("ДЕНЬ РОЖДЕНИЯ - ", prettier_check_tool(self.birth_day))

        change_row(tables, 28,0,4,5, self.birth_month, self.font_settings)
        print("МЕСЯЦ РОЖДЕНИЯ - ",prettier_check_tool(self.birth_month))

        change_row(tables, 28,0,7,10, self.birth_year, self.font_settings)
        print("ГОД РОЖДЕНИЯ - ", prettier_check_tool(self.birth_year))

        if len(self.passport_series) == 1:
            change_row(tables, 30,0, 6,6, " ", self.font_settings)
            change_row(tables, 30,0, 7,7, self.passport_series, self.font_settings)
        elif len(self.passport_series) == 2:
            change_row(tables, 30,0, 6,7, self.passport_series, self.font_settings)
        else:
            print("___ОПЕЧАТКА?___")
        print("СЕРИЯ ПАСПОРТА - ", prettier_check_tool(self.passport_series))

        change_row(tables, 30,0,9,17, self.passport_number, self.font_settings)
        print("НОМЕР ПАСПОРТА - ", prettier_check_tool(self.passport_number))

        change_row(tables, 30,0,19,20, self.passport_date_day, self.font_settings)
        print("ДЕНЬ ВЫДАЧИ ПАСПОРТА - ", prettier_check_tool(self.passport_date_day))

        change_row(tables, 30,0,22,23, self.passport_date_month, self.font_settings)
        print("МЕСЯЦ ВЫДАЧИ ПАСПОРТА - ", prettier_check_tool(self.passport_date_month))

        change_row(tables, 30,0,25,28, self.passport_date_year, self.font_settings)
        print("ГОД ВЫДАЧИ ПАСПОРТА - ", prettier_check_tool(self.passport_date_year))

        change_row(tables, 43,0,0,33, self.profession_name, self.font_settings)
        print("ПРОФЕССИЯ - ", prettier_check_tool(self.profession_name))

        self.save_document(self.document, self.name, self.lastname)

