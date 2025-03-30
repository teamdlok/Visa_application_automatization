from pathlib import Path
from docx import Document
import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import calendar
from FUNCTIONS_AND_CLASSES.FUNCTIONS import *


company_pathes = {
    "АВАНТА": "/templates/АВАНТА/ЗАЯВЛЕНИЕ/zayavlenie_template.docx",
    "ПРОФЕССИОНАЛ": "/templates/ПРОФЕССИОНАЛ/ЗАЯВЛЕНИЕ/zayavlenie_template.docx",
    "ВИЗАР ВОСТОК": "/templates/ВИЗАР ВОСТОК/ЗАЯВЛЕНИЕ/zayavlenie_template.docx",
    "ТЭНФЭЙ": "/templates/ТЭНФЭЙ/ЗАЯВЛЕНИЕ/zayavlenie_template.docx",
}

font_params = {
    "fontname": "Arial Narrow",
    "fontsize": 11,
    "bold": True,
}

company_choice = int(input("ВЫБЕРИТЕ КОМПАНИЮ \n 1 - АВАНТА \n 2 - ПРОФЕССИОНАЛ \n 3 - ВИЗАР ВОСТОК \n 4 - ТЭНФЭЙ \n"))
if company_choice == 1:
    company_name = "АВАНТА"
    company_path = company_pathes[company_name]
elif company_choice == 2:
    company_name = "ПРОФЕССИОНАЛ"
    company_path = company_pathes[company_name]
elif company_choice == 3:
    company_name = "ВИЗАР ВОСТОК"
    company_path = company_pathes[company_name]
elif company_choice == 4:
    company_name = "ТЭНФЭЙ"
    company_path = company_pathes[company_name]
    

month_dict = {}

wordDoc = Document(f'C:/Users/ekole/Desktop/test/{company_path}')
tables = wordDoc.tables


# Таблица 0 багованная, потому для неё нужен отдельный функционал
def change_row_durdom(tables, table_number, row_number, from_indx, to_indx, text, font_params, igor_toxic = False):
    text = list(text)
    text_is_bigger = False
    counter = 0
    
    row = tables[table_number].rows[row_number]

    for index, cell in enumerate(row.cells):
        try:
            if not index % 2 == 0:

                if to_indx - from_indx + 1 < len(text):
                    text_is_bigger = True
                if index >= from_indx and index <= to_indx:
                    if igor_toxic:
                        cell.text = text
                    else:
                        cell.text = text.pop(0)
                    counter = counter + 1
                    # print(f" символ - {cell.text} счетчик - {counter}")

                    change_cell_font(cell, **font_params)
                    paragraphs = cell.paragraphs
                    paragraphs
                    for paragraph in paragraphs:
                        if "align" in font_params:
                            if font_params['align'] == "LEFT":
                                paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
                
                elif to_indx - from_indx + 1 < len(text): # Тут и скрывается проблема. не подпадая под первое условие, оно подпадает под это, и меняет строку
                    print(f"{to_indx - from_indx} меньше чем {len(text)}")

        except IndexError:
            cell.text = ""
            # print("<<ERROR>>")
            pass
        
    if text_is_bigger:
        to_index_len = len(tables[table_number].rows[row_number].cells)
        if len(tables[table_number].rows) > 1:
            change_row(tables=tables, table_number=table_number, row_number=row_number+1,from_indx=0, to_indx=to_index_len, text=text, font_params=font_params)
        else:
            change_row(tables=tables, table_number=table_number+1, row_number=row_number,from_indx=0, to_indx=to_index_len, text=text, font_params=font_params)

rnp_input = input_prettier("номер рнп")
if rnp_input == []:
    rnp_input = "РНП № 2400"
elif len(rnp_input) > 6:
    rnp_input = f"РНП № {''.join(rnp_input)}"
else:
    rnp_input = f"РНП № {''.join(rnp_input)}"
    print(f"_________Ошибка? {''.join(rnp_input)}?")
change_row_durdom(tables, 0,0,0,49, rnp_input, dict(fontname = "Times New Roman" , fontsize = 12, bold = True, align = "LEFT"), igor_toxic=True)
print(prettier_check_tool(rnp_input))

day_rnp_start = input_prettier("день начала разрешения", True)
change_row_durdom(tables, 0, 2, 4, 7, day_rnp_start, font_params)
change_row(tables, 38,0,1,2, day_rnp_start, font_params)
print(prettier_check_tool(day_rnp_start))

month_rnp_start = input_prettier("месяц начала разрешения", True)
change_row_durdom(tables, 0, 2, 10, 13, month_rnp_start, font_params)
change_row(tables, 38,0,4,5, month_rnp_start, font_params)
print(prettier_check_tool(month_rnp_start))

year_rnp_start = input_prettier("год начала разрешения", False, True)
change_row_durdom(tables, 0, 2, 16, 23, year_rnp_start, font_params)
change_row(tables, 38,0,7,10, year_rnp_start, font_params)
print(prettier_check_tool(year_rnp_start))

day_rnp_end = input_prettier("день конца разрешения", True)
change_row_durdom(tables, 0, 2, 26, 28, day_rnp_end, font_params)
change_row(tables, 0,2,26,26, day_rnp_end, font_params)
change_row(tables, 38,0,12,13, day_rnp_end, font_params)
print(prettier_check_tool(day_rnp_end))

month_rnp_end = input_prettier("месяц конца разрешения", True)
change_row_durdom(tables, 0, 2, 31, 33, month_rnp_end, font_params)
change_row(tables, 38,0,15,16, month_rnp_end, font_params)
print(prettier_check_tool(month_rnp_end))

year_rnp_end = input_prettier("год конца разрешения", False, True)
change_row_durdom(tables, 0, 2, 37, 43, year_rnp_end, font_params)
change_row(tables, 38,0,18,21, year_rnp_end, font_params)
print(prettier_check_tool(year_rnp_end))

lastname_input = input_prettier("ФАМИЛИЮ")
change_row_durdom(tables, 0, 4, 6, 55, lastname_input, font_params)
print(prettier_check_tool(lastname_input))

name_input = input_prettier("ИМЯ")
change_row(tables, 1,0,5,20, name_input, font_params)
print(prettier_check_tool(name_input))

# namechange_input = input_prettier("СВЕДЕНИЯ ОБ ИЗМЕНЕНИИ ФИО")
# if str(''.join(namechange_input)) == "":
#     namechange_input = list("НЕ МЕНЯЛ")
#     namechange_check = False
# else:
#     namechange_check = True
# change_row(tables, 3,0,1,24, namechange_input, font_params)
# print(prettier_check_tool(namechange_input))

# birthplace_input = input_prettier("МЕСТО РОЖДЕНИЯ")
# change_row(tables, 8,0,15,28, birthplace_input, font_params)
# print(prettier_check_tool(birthplace_input))

# birth_day = input_prettier("ДЕНЬ РОЖДЕНИЯ", True)
# change_row(tables, 11,0,1,2, birth_day, font_params)
# print(prettier_check_tool(birth_day))

# birth_month = input_prettier("МЕСЯЦ РОЖДЕНИЯ (слово, например AUG)")
# month_number = numbered_month(''.join(birth_month))
# if month_number < 10:
#     month_number = f"0{month_number}"
# month_number = list(str(month_number))
# change_row(tables, 11,0,4,5, month_number, font_params)
# print(prettier_check_tool(month_number))

# birth_year = input_prettier("ГОД РОЖДЕНИЯ", False, True)
# birth_year_int = int(''.join(birth_year))
# if birth_year_int > 2005:
#     print(f"______Ошибка? Родился в {birth_year_int}?______")
# change_row(tables, 11,0,7,10, birth_year, font_params)
# print(prettier_check_tool(birth_year))

# sex = input_prettier("ПОЛ (М) - (Ж)")
# if str(''.join(sex)) == "М":
#     change_row(tables, 11,0, 12,12, sex, font_params)
#     change_row(tables, 11,0, 14,14, "", font_params)
# elif str(''.join(sex)) == "Ж":
#     change_row(tables, 11,0, 12,12, "", font_params)
#     change_row(tables, 11,0, 14,14, sex, font_params)
#     if not namechange_check:
#         change_row(tables, 3,0,9,9, "А", font_params)

# else:
#     print("МЫ ТУТ?")
# print(prettier_check_tool(sex))

# passport_series = input_prettier("СЕРИЮ ПАСПОРТА")
# if len(passport_series) == 1:
#     change_row(tables, 13, 0, 6, 6, " ", font_params)
#     change_row(tables, 13, 0, 7, 7, passport_series, font_params)
# elif len(passport_series) == 2:
#     change_row(tables, 13, 0, 6, 7, passport_series, font_params)
# else:
#     print("___ОПЕЧАТКА?___")
# print(prettier_check_tool(passport_series))

# passport_number = input_prettier("НОМЕР ПАСПОРТА")
# if len(passport_number) < 6:
#     print(f"_____Ошибка? В номере паспорта всего {len(passport_number)} цифры")
# change_row(tables, 13,0,9,16, passport_number, font_params)
# print(prettier_check_tool(passport_number))

# passport_date_day = input_prettier("ДЕНЬ ВЫДАЧИ ПАСПОРТА", True)
# change_row(tables, 13,0,18,19, passport_date_day, font_params)
# print(prettier_check_tool(passport_date_day))

# passport_date_month = input_prettier("МЕСЯЦ ВЫДАЧИ ПАСПОРТА", True)
# change_row(tables, 13,0,21,22, passport_date_month, font_params)
# print(prettier_check_tool(passport_date_month))

# passport_date_year = input_prettier("ГОД ВЫДАЧИ ПАСПОРТА", False, True)
# passport_year_int = int(''.join(passport_date_year))
# if passport_year_int < 2014:
#     print(f"___Ошибка? Паспорт сделан в {passport_year_int}?___")
# change_row(tables, 13,0,24,27, passport_date_year, font_params)
# print(prettier_check_tool(passport_date_year))

# passport_creator = input_prettier("КЕМ ВЫДАН ПАСПОРТ")
# change_row(tables, 14,0,1,31, passport_creator, font_params)
# print(prettier_check_tool(passport_creator))

# adress_of_living = input_prettier("АДРЕС ПРОПИСКИ")
# change_row(tables, 17,0,1,21, adress_of_living, font_params)
# print(prettier_check_tool(adress_of_living))

# profession_name = input_prettier("ПРОФЕССИЮ")
# change_row(tables, 20,0,1,22, profession_name, font_params)
# print(prettier_check_tool(profession_name))


Path(f"OUTPUT/REQUEST_FOR_VISA/{company_name}").mkdir(parents=False, exist_ok=True)
current_time = datetime.now().strftime('%d.%m.%Y')

wordDoc.save(f'OUTPUT/REQUEST_FOR_VISA/{company_name}/{''.join(lastname_input)} {''.join(name_input)} - {current_time} .docx')
